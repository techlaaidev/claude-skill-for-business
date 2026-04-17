"""
Builder báo cáo Research sản phẩm mới — output .md + .docx song song.
Author: Techla — v1.0.0 — License: xem LICENSE.md

Usage:
    python scripts/build_report.py --data research.json --config config.yaml \
        --md research.md --docx research.docx

data.json schema (do Claude tạo SAU khi đã research xong bằng WebSearch/WebFetch):
{
    "topic": "Trend matcha 2026",
    "scope": "both",                           # vn | intl | both
    "generated_at": "2026-04-17",
    "summary": "Tóm tắt 3-5 câu về trend matcha 2026...",
    "trends": [                                # top 5 xu hướng chính
        {
            "title": "Matcha đen (roasted matcha)",
            "description": "Xu hướng matcha rang lên để giảm vị chát...",
            "examples": ["Starbucks Japan", "Blue Bottle"],
            "source_refs": [1, 3]              # index vào mảng sources
        }
    ],
    "vn_intl_analysis": "So sánh VN (giá 45-65k) vs quốc tế (giá 120-180k)...",
    "case_studies": [
        {
            "brand": "Phúc Long",
            "highlight": "Ra mắt matcha latte đá xay 2026-02, doanh thu tăng 15%",
            "learning": "Kết hợp với kem cheese, topping đậu đỏ",
            "source_refs": [2]
        }
    ],
    "opportunities": [                          # cơ hội cho quán tầm trung VN
        "Thêm menu matcha latte nóng + đá ở 3 mức đường (ít/vừa/nhiều)",
        "Bán kèm bánh matcha cookies giá 25k để tăng ticket size"
    ],
    "sources": [                                # danh sách nguồn đầy đủ (cite 1-indexed trong báo cáo)
        {
            "title": "Matcha Market Trends 2026",
            "url": "https://example.com/...",
            "site": "Perfect Daily Grind",
            "date": "2026-03-15",
            "quote": "Matcha consumption in Asia grew 24% YoY"  # optional, ≤15 từ
        }
    ]
}
"""
from __future__ import annotations
import os
import sys
import json
import argparse
from pathlib import Path
from datetime import date as _date

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import yaml
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


BASE_FONT = "Calibri"
BASE_SIZE = 11
COLOR_TITLE = RGBColor(0x1F, 0x4E, 0x79)
COLOR_H2 = RGBColor(0x2E, 0x75, 0xB6)
COLOR_GRAY = RGBColor(0x7F, 0x8C, 0x8D)

FOOTER_TEXT = "Báo cáo tạo bởi Techla — Skill Research sản phẩm mới v1.0.0"


# ═══════════════════════════════════════════════════════════════════
# MARKDOWN OUTPUT
# ═══════════════════════════════════════════════════════════════════

def _cite(refs: list[int]) -> str:
    if not refs:
        return ""
    return " " + "".join(f"[^{r}]" for r in refs)


def build_markdown(data: dict, config: dict) -> str:
    topic = data.get("topic", "(chưa có chủ đề)")
    scope = data.get("scope", "both")
    scope_label = {"vn": "Thị trường Việt Nam", "intl": "Thị trường quốc tế",
                   "both": "VN + Quốc tế"}.get(scope, scope)
    generated = data.get("generated_at") or _date.today().isoformat()

    lines = []
    lines.append(f"# Báo cáo Research: {topic}")
    lines.append("")
    lines.append(f"**Phạm vi:** {scope_label}  ")
    lines.append(f"**Ngày tổng hợp:** {generated}  ")
    lines.append(f"**Người tổng hợp:** Claude (skill Techla Research sản phẩm mới)")
    lines.append("")
    lines.append("---")
    lines.append("")

    # Tóm tắt
    lines.append("## 1. Tóm tắt")
    lines.append("")
    lines.append(data.get("summary", "(chưa có tóm tắt)"))
    lines.append("")

    # Xu hướng chính
    trends = data.get("trends") or []
    lines.append(f"## 2. Xu hướng chính (top {len(trends)})")
    lines.append("")
    for i, t in enumerate(trends, 1):
        cite = _cite(t.get("source_refs") or [])
        lines.append(f"### 2.{i}. {t.get('title', '(chưa đặt tên)')}")
        lines.append("")
        lines.append(f"{t.get('description', '')}{cite}")
        lines.append("")
        examples = t.get("examples") or []
        if examples:
            lines.append(f"**Ví dụ:** {', '.join(examples)}")
            lines.append("")

    # Phân tích VN vs quốc tế
    analysis = data.get("vn_intl_analysis", "").strip()
    if analysis:
        lines.append("## 3. Phân tích VN vs Quốc tế")
        lines.append("")
        lines.append(analysis)
        lines.append("")

    # Case studies
    cases = data.get("case_studies") or []
    if cases:
        lines.append("## 4. Case study nổi bật")
        lines.append("")
        for i, c in enumerate(cases, 1):
            cite = _cite(c.get("source_refs") or [])
            lines.append(f"### 4.{i}. {c.get('brand', '(thương hiệu)')}")
            lines.append("")
            lines.append(f"**Điểm nổi bật:** {c.get('highlight', '')}{cite}")
            lines.append("")
            learn = c.get("learning", "").strip()
            if learn:
                lines.append(f"**Bài học áp dụng:** {learn}")
                lines.append("")

    # Cơ hội áp dụng
    opps = data.get("opportunities") or []
    if opps:
        lines.append("## 5. Cơ hội áp dụng cho quán cafe tầm trung VN")
        lines.append("")
        for i, o in enumerate(opps, 1):
            lines.append(f"{i}. {o}")
        lines.append("")

    # Nguồn
    sources = data.get("sources") or []
    if sources:
        lines.append("## 6. Nguồn tham khảo")
        lines.append("")
        for i, s in enumerate(sources, 1):
            title = s.get("title", "(không có tiêu đề)")
            url = s.get("url", "")
            site = s.get("site", "")
            dt = s.get("date", "")
            meta = " — ".join([x for x in [site, dt] if x])
            if url:
                lines.append(f"{i}. [{title}]({url}){(' — ' + meta) if meta else ''}")
            else:
                lines.append(f"{i}. {title}{(' — ' + meta) if meta else ''}")
            quote = s.get("quote", "").strip()
            if quote:
                lines.append(f"   > \"{quote}\"")
        lines.append("")

    lines.append("---")
    lines.append("")
    lines.append(f"*{FOOTER_TEXT}*")
    lines.append("")

    return "\n".join(lines)


# ═══════════════════════════════════════════════════════════════════
# DOCX OUTPUT
# ═══════════════════════════════════════════════════════════════════

def _set_run(run, *, size=BASE_SIZE, bold=False, italic=False, color=None):
    run.font.name = BASE_FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def _add_para(doc, text, *, size=BASE_SIZE, bold=False, italic=False,
              color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    _set_run(r, size=size, bold=bold, italic=italic, color=color)
    return p


def _add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    _set_run(r, size=18, bold=True, color=COLOR_TITLE)
    return p


def _add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    _set_run(r, size=14, bold=True, color=COLOR_H2)
    return p


def _add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    _set_run(r, size=12, bold=True, color=COLOR_H2)
    return p


def _add_meta_line(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{label}: ")
    _set_run(r1, bold=True)
    r2 = p.add_run(value)
    _set_run(r2)


def _add_cite_superscript(paragraph, refs):
    if not refs:
        return
    for r in refs:
        run = paragraph.add_run(f"[{r}]")
        _set_run(run, size=BASE_SIZE - 2, color=COLOR_H2)
        run.font.superscript = True


def build_docx(data: dict, config: dict) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    style = doc.styles["Normal"]
    style.font.name = BASE_FONT
    style.font.size = Pt(BASE_SIZE)

    topic = data.get("topic", "(chưa có chủ đề)")
    scope = data.get("scope", "both")
    scope_label = {"vn": "Thị trường Việt Nam", "intl": "Thị trường quốc tế",
                   "both": "VN + Quốc tế"}.get(scope, scope)
    generated = data.get("generated_at") or _date.today().isoformat()

    _add_h1(doc, f"Báo cáo Research: {topic}")
    _add_meta_line(doc, "Phạm vi", scope_label)
    _add_meta_line(doc, "Ngày tổng hợp", generated)
    _add_meta_line(doc, "Người tổng hợp", "Claude (skill Techla Research sản phẩm mới)")
    _add_para(doc, "", space_after=6)

    # 1. Tóm tắt
    _add_h2(doc, "1. Tóm tắt")
    _add_para(doc, data.get("summary", "(chưa có tóm tắt)"))

    # 2. Xu hướng
    trends = data.get("trends") or []
    _add_h2(doc, f"2. Xu hướng chính (top {len(trends)})")
    for i, t in enumerate(trends, 1):
        _add_h3(doc, f"2.{i}. {t.get('title', '(chưa đặt tên)')}")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(t.get("description", ""))
        _set_run(r)
        _add_cite_superscript(p, t.get("source_refs") or [])

        examples = t.get("examples") or []
        if examples:
            p2 = doc.add_paragraph()
            r1 = p2.add_run("Ví dụ: ")
            _set_run(r1, bold=True)
            r2 = p2.add_run(", ".join(examples))
            _set_run(r2)

    # 3. Phân tích VN vs quốc tế
    analysis = data.get("vn_intl_analysis", "").strip()
    if analysis:
        _add_h2(doc, "3. Phân tích VN vs Quốc tế")
        _add_para(doc, analysis)

    # 4. Case studies
    cases = data.get("case_studies") or []
    if cases:
        _add_h2(doc, "4. Case study nổi bật")
        for i, c in enumerate(cases, 1):
            _add_h3(doc, f"4.{i}. {c.get('brand', '(thương hiệu)')}")
            p = doc.add_paragraph()
            r1 = p.add_run("Điểm nổi bật: ")
            _set_run(r1, bold=True)
            r2 = p.add_run(c.get("highlight", ""))
            _set_run(r2)
            _add_cite_superscript(p, c.get("source_refs") or [])

            learn = c.get("learning", "").strip()
            if learn:
                p2 = doc.add_paragraph()
                r1 = p2.add_run("Bài học áp dụng: ")
                _set_run(r1, bold=True)
                r2 = p2.add_run(learn)
                _set_run(r2)

    # 5. Cơ hội
    opps = data.get("opportunities") or []
    if opps:
        _add_h2(doc, "5. Cơ hội áp dụng cho quán cafe tầm trung VN")
        for o in opps:
            p = doc.add_paragraph(style="List Number")
            r = p.add_run(o)
            _set_run(r)

    # 6. Nguồn
    sources = data.get("sources") or []
    if sources:
        _add_h2(doc, "6. Nguồn tham khảo")
        for i, s in enumerate(sources, 1):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r0 = p.add_run(f"[{i}] ")
            _set_run(r0, bold=True, color=COLOR_H2)
            r1 = p.add_run(s.get("title", "(không có tiêu đề)"))
            _set_run(r1, bold=True)
            url = s.get("url", "")
            if url:
                r2 = p.add_run(f" — {url}")
                _set_run(r2, size=BASE_SIZE - 1, color=COLOR_GRAY)
            meta_parts = [s.get("site", ""), s.get("date", "")]
            meta = " — ".join([x for x in meta_parts if x])
            if meta:
                r3 = p.add_run(f" ({meta})")
                _set_run(r3, italic=True, size=BASE_SIZE - 1, color=COLOR_GRAY)
            quote = s.get("quote", "").strip()
            if quote:
                pq = doc.add_paragraph()
                pq.paragraph_format.left_indent = Cm(1.0)
                pq.paragraph_format.space_after = Pt(6)
                rq = pq.add_run(f"\u201C{quote}\u201D")
                _set_run(rq, italic=True, color=COLOR_GRAY)

    # Footer
    doc.add_paragraph()
    pf = doc.add_paragraph()
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.paragraph_format.space_before = Pt(12)
    rf = pf.add_run(FOOTER_TEXT)
    _set_run(rf, size=9, italic=True, color=COLOR_GRAY)

    return doc


# ═══════════════════════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════════════════════

def main():
    ap = argparse.ArgumentParser(description="Builder báo cáo Research sản phẩm mới.")
    ap.add_argument("--data", required=True, help="File JSON chứa kết quả research.")
    ap.add_argument("--config", default=None)
    ap.add_argument("--md", default=None, help="Output Markdown path.")
    ap.add_argument("--docx", default=None, help="Output DOCX path.")
    args = ap.parse_args()

    if not args.md and not args.docx:
        print("Lỗi: phải có ít nhất --md hoặc --docx.", file=sys.stderr); sys.exit(2)

    if not Path(args.data).exists():
        print(f"Lỗi: không tìm thấy '{args.data}'.", file=sys.stderr); sys.exit(1)
    with open(args.data, "r", encoding="utf-8") as f:
        data = json.load(f)

    config = {}
    if args.config and Path(args.config).exists():
        with open(args.config, "r", encoding="utf-8") as f:
            config = yaml.safe_load(f) or {}

    if not data.get("topic"):
        print("Lỗi: thiếu 'topic' trong data.", file=sys.stderr); sys.exit(3)

    if args.md:
        md_text = build_markdown(data, config)
        Path(args.md).parent.mkdir(parents=True, exist_ok=True)
        with open(args.md, "w", encoding="utf-8") as f:
            f.write(md_text)
        print(f"OK MD: {args.md}")

    if args.docx:
        doc = build_docx(data, config)
        Path(args.docx).parent.mkdir(parents=True, exist_ok=True)
        doc.save(args.docx)
        print(f"OK DOCX: {args.docx}")


if __name__ == "__main__":
    main()
