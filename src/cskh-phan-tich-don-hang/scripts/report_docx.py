"""
Build báo cáo .docx từ analysis.json.
Author: Techla — v1.0.0 — License: xem LICENSE.md

Bố cục: A4, margin 2cm, Arial 11pt, 11 section theo spec.
"""
from __future__ import annotations
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from vn_format import format_money, format_int, format_float, format_percent


COLOR_PRIMARY = RGBColor(0x1F, 0x4E, 0x79)
COLOR_HEADER = "2E75B6"
COLOR_RED = RGBColor(0xC0, 0x39, 0x2B)
COLOR_ORANGE = RGBColor(0xE6, 0x7E, 0x22)
COLOR_GREEN = RGBColor(0x27, 0xAE, 0x60)
COLOR_GRAY = RGBColor(0x7F, 0x8C, 0x8D)

BASE_FONT = "Arial"


def _set_font(run, *, size=11, bold=False, color=None, italic=False):
    run.font.name = BASE_FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    # East Asian font fallback
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), BASE_FONT)


def _shade_cell(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _setup_doc(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    # Default style
    style = doc.styles["Normal"]
    style.font.name = BASE_FONT
    style.font.size = Pt(11)


def _add_heading(doc, text, level=1, color=COLOR_PRIMARY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    size = {1: 16, 2: 13, 3: 11}.get(level, 11)
    _set_font(run, size=size, bold=True, color=color)
    return p


def _add_table(doc, headers: list[str], rows: list[list[str]], widths_cm: list[float] | None = None):
    if not headers:
        return None
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(str(h))
        _set_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _shade_cell(hdr[i], COLOR_HEADER)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row in rows:
        tr = table.add_row().cells
        for i, val in enumerate(row):
            tr[i].text = ""
            p = tr[i].paragraphs[0]
            run = p.add_run(str(val))
            _set_font(run, size=10)

    if widths_cm:
        for row in table.rows:
            for i, w in enumerate(widths_cm):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    # Post-spacing
    doc.add_paragraph()
    return table


def build_docx(a: dict, output_path: str):
    doc = Document()
    _setup_doc(doc)

    # -------- Title page --------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(120)
    title_p.paragraph_format.space_after = Pt(12)
    run = title_p.add_run("BÁO CÁO PHÂN TÍCH ĐƠN HÀNG\n& CHĂM SÓC KHÁCH HÀNG")
    _set_font(run, size=22, bold=True, color=COLOR_PRIMARY)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(sub_p.add_run(f"Nguồn: {a.get('source_file', '-')}"), size=11, italic=True, color=COLOR_GRAY)
    sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(sub2.add_run(f"Tạo lúc: {a['generated_at']} · Ngày tham chiếu: {a['reference_date']}"),
              size=11, italic=True, color=COLOR_GRAY)

    doc.add_page_break()

    # -------- Sections --------
    _sec_overview(doc, a)
    _sec_churn(doc, a)
    _sec_vip(doc, a)
    _sec_cross_sell(doc, a)
    _sec_star(doc, a)
    _sec_dead(doc, a)
    _sec_mom(doc, a)
    _sec_monthly(doc, a)
    _sec_decreasing(doc, a)
    if a["config_used"].get("enable_basket_analysis", True):
        _sec_basket(doc, a)
    _sec_actions(doc, a)

    # Footer cuối
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(foot.add_run("Báo cáo tạo bởi Techla — Skill CSKH Phân tích đơn hàng v1.0.0"),
              size=9, italic=True, color=COLOR_GRAY)

    doc.save(output_path)


def _sec_overview(doc, a):
    ov = a["overview"]
    _add_heading(doc, "1. Tóm tắt điều hành")
    rows = [
        ["Khoảng thời gian", f"{ov['period_start']} → {ov['period_end']}"],
        ["Tổng khách hàng", format_int(ov["total_customers"])],
        ["Tổng sản phẩm", format_int(ov["total_products"])],
        ["Tổng đơn hàng", format_int(ov["total_orders"])],
        ["Tổng số lượng", format_float(ov["total_qty"])],
        ["Doanh thu ước tính", format_money(ov["total_revenue"])],
    ]
    _add_table(doc, ["Chỉ số", "Giá trị"], rows, widths_cm=[6, 10])


def _sec_churn(doc, a):
    rows_data = a["churn_warnings"]
    _add_heading(doc, "2. ⚠️ Khách sắp mất", color=COLOR_RED)
    if not rows_data:
        doc.add_paragraph("Không có khách hàng nào có dấu hiệu churn.")
        return
    rows = []
    for r in rows_data:
        hint = (r.get("care_history") or r.get("care_goal") or "—")
        rows.append([
            r["name"], r.get("phone") or "—", r["last_order_date"],
            f"{r['days_since_last_order']} ngày", f"{format_float(r['avg_gap_days'])} ngày",
            format_money(r["total_revenue"]), hint[:80],
        ])
    _add_table(doc, ["Khách", "SĐT", "Đơn cuối", "Đã lâu", "Nhịp TB", "Doanh thu", "Gợi ý"], rows)


def _sec_vip(doc, a):
    rows_data = a["top_vip"]
    _add_heading(doc, "3. 👑 Top VIP")
    if not rows_data:
        doc.add_paragraph("Không có dữ liệu VIP."); return
    rows = []
    for i, r in enumerate(rows_data, 1):
        rows.append([
            str(i), r["name"], r.get("phone") or "—",
            format_money(r["total_revenue"]), format_int(r["total_orders"]),
            format_float(r["total_qty"]),
        ])
    _add_table(doc, ["#", "Khách", "SĐT", "Doanh thu", "Đơn", "Số lượng"], rows)


def _sec_cross_sell(doc, a):
    rows = a["cross_sell_opportunities"]
    _add_heading(doc, "4. 💡 Cơ hội bán thêm")
    if not rows:
        doc.add_paragraph("Chưa có ghi chú cơ hội nào."); return
    for r in rows:
        p = doc.add_paragraph()
        run = p.add_run(f"• {r['name']} — {r.get('phone') or '—'}")
        _set_font(run, size=11, bold=True)
        if r.get("opportunity"):
            pi = doc.add_paragraph(); pi.paragraph_format.left_indent = Cm(0.8)
            _set_font(pi.add_run(f"Cơ hội: {r['opportunity']}"), size=10)
        if r.get("action_needed"):
            pi = doc.add_paragraph(); pi.paragraph_format.left_indent = Cm(0.8)
            _set_font(pi.add_run(f"Cần làm: {r['action_needed']}"), size=10)


def _sec_star(doc, a):
    rows_data = a["star_products"]
    _add_heading(doc, "5. ⭐ Sản phẩm ngôi sao")
    if not rows_data:
        doc.add_paragraph("Không có dữ liệu."); return
    rows = []
    for i, r in enumerate(rows_data, 1):
        rows.append([
            str(i), r["name"], format_money(r["revenue"]),
            format_int(r["orders"]), format_int(r["unique_customers"]), format_float(r["qty"]),
        ])
    _add_table(doc, ["#", "SP", "Doanh thu", "Đơn", "KH mua", "SL"], rows)


def _sec_dead(doc, a):
    rows_data = a["dead_products"]
    _add_heading(doc, "6. 💀 Sản phẩm xác sống")
    if not rows_data:
        doc.add_paragraph("Không có sản phẩm xác sống."); return
    rows = []
    for r in rows_data:
        rows.append([r["name"], r["last_order_date"] or "—", r["reason"]])
    _add_table(doc, ["SP", "Đơn cuối", "Lý do"], rows)


def _sec_mom(doc, a):
    mom = a["mom_compare"]
    _add_heading(doc, "7. 📊 So sánh MoM")
    if not mom:
        doc.add_paragraph("Chỉ có dữ liệu 1 tháng — không tính được MoM."); return
    p = doc.add_paragraph()
    _set_font(p.add_run(f"So sánh {mom['current_month']} với {mom['previous_month']}:"), size=11, bold=True)
    items = [
        ("Đơn hàng", mom["orders_change_pct"]),
        ("Doanh thu", mom["revenue_change_pct"]),
        ("Khách", mom["customers_change_pct"]),
        ("Số lượng", mom["qty_change_pct"]),
    ]
    for label, val in items:
        pi = doc.add_paragraph(style="List Bullet")
        _set_font(pi.add_run(f"{label}: {format_percent(val)}"), size=11)


def _sec_monthly(doc, a):
    rows_data = a["monthly"]
    _add_heading(doc, "8. 📅 Chi tiết theo tháng")
    if not rows_data:
        doc.add_paragraph("Không có dữ liệu."); return
    rows = []
    for r in rows_data:
        rows.append([
            r["month"], format_int(r["orders"]), format_float(r["qty"]),
            format_money(r["revenue"]), format_int(r["unique_customers"]),
        ])
    _add_table(doc, ["Tháng", "Đơn", "Số lượng", "Doanh thu", "KH duy nhất"], rows)


def _sec_decreasing(doc, a):
    rows_data = a["decreasing_rhythm"]
    _add_heading(doc, "9. 📉 Khách giảm nhịp")
    if not rows_data:
        doc.add_paragraph("Không có khách nào giảm nhịp đột ngột."); return
    rows = []
    for r in rows_data:
        rows.append([
            r["name"], r.get("phone") or "—", r["last_month"],
            format_int(r["last_month_orders"]), format_float(r["prev_avg_orders_per_month"]),
            f"-{format_float(r['drop_pct'])}%", format_money(r["total_revenue"]),
        ])
    _add_table(doc, ["Khách", "SĐT", "Tháng cuối", "Đơn cuối", "TB trước", "% giảm", "Doanh thu"], rows)


def _sec_basket(doc, a):
    rows_data = a["basket_pairs"]
    _add_heading(doc, "10. 🛒 Cặp SP hay mua cùng")
    if not rows_data:
        doc.add_paragraph("Không tìm thấy cặp SP mua chung."); return
    rows = []
    for i, r in enumerate(rows_data, 1):
        rows.append([str(i), r["product_a_name"], r["product_b_name"], format_int(r["count"])])
    _add_table(doc, ["#", "SP A", "SP B", "Số lần mua chung"], rows)


def _sec_actions(doc, a):
    rows = a["action_list"]
    _add_heading(doc, "11. 🎯 Action list")
    if not rows:
        doc.add_paragraph("Không có action cần làm."); return
    color_map = {"CAO": COLOR_RED, "TRUNG BÌNH": COLOR_ORANGE, "THẤP": COLOR_GREEN}
    for r in rows:
        p = doc.add_paragraph()
        badge = p.add_run(f"[{r['priority']}] ")
        _set_font(badge, size=11, bold=True, color=color_map.get(r["priority"], COLOR_GRAY))
        body = p.add_run(f"{r['type']} — {r['customer']} ({r.get('phone') or '—'})")
        _set_font(body, size=11, bold=True)
        detail = doc.add_paragraph(); detail.paragraph_format.left_indent = Cm(0.8)
        _set_font(detail.add_run(r["detail"]), size=10)
        if r.get("care_hint"):
            hint = doc.add_paragraph(); hint.paragraph_format.left_indent = Cm(0.8)
            _set_font(hint.add_run(f"Gợi ý: {r['care_hint']}"), size=10, italic=True, color=COLOR_GRAY)
