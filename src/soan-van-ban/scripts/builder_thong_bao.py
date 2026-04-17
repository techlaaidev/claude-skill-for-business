"""
Builder: Thông báo nội bộ quán cafe.
Author: Techla — v1.0.0 — License: xem LICENSE.md

Usage:
    python scripts/builder_thong_bao.py --data data.json --config config.yaml --output thong-bao.docx

data.json schema:
{
    "subject": "Thay đổi lịch ca tuần 20/04 – 26/04",
    "category": "doi-ca",                  # doi-ca | mon-moi | quy-dinh | nhac-nho | nghi-le | khac
    "recipient": "toàn thể nhân viên",     # hoặc "ca sáng", "team pha chế"...
    "content": "Từ tuần sau...",           # nội dung chính (có thể nhiều đoạn, ngăn bởi \n\n)
    "bullets": ["..."],                    # optional — danh sách gạch đầu dòng
    "effective_from": "2026-04-20",        # optional
    "effective_to": "2026-04-26",          # optional
    "signed_by": "Trần Thị Chủ",
    "signed_title": "Chủ quán",
    "signed_date": "2026-04-17",
    "location": "Hà Nội"
}
"""
from __future__ import annotations
import os
import sys
import json
import argparse
from pathlib import Path

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import yaml

from doc_style import (
    new_doc, add_centered, add_paragraph_styled, add_footer,
    vn_date, vn_header_quoc_hieu, set_run_font, BASE_SIZE
)


CATEGORY_LABELS = {
    "doi-ca": "V/v thay đổi lịch ca",
    "mon-moi": "V/v ra món mới / đổi công thức",
    "quy-dinh": "V/v áp dụng quy định mới",
    "nhac-nho": "V/v nhắc nhở quy định nội bộ",
    "nghi-le": "V/v lịch nghỉ lễ",
    "khac": "V/v thông báo nội bộ",
}


def build_announcement(data: dict, config: dict) -> "Document":
    ten_quan = config.get("ten_quan", "Quán Cafe")

    subject = data.get("subject", "Thông báo nội bộ")
    category = data.get("category", "khac").lower()
    category_label = CATEGORY_LABELS.get(category, CATEGORY_LABELS["khac"])
    recipient = data.get("recipient", "toàn thể nhân viên")
    content = data.get("content", "").strip()
    bullets = data.get("bullets") or []
    location = data.get("location", "")
    signed_by = data.get("signed_by", "(Chưa có)")
    signed_title = data.get("signed_title", "Chủ quán")
    signed_date = data.get("signed_date") or data.get("effective_from")

    doc = new_doc()
    vn_header_quoc_hieu(doc)

    # Tiêu đề
    add_centered(doc, "THÔNG BÁO", bold=True, size=16, space_before=6, space_after=2)
    add_centered(doc, category_label, italic=True, size=12, space_after=12)

    # Body
    add_paragraph_styled(doc, f"Kính gửi: {recipient.capitalize() if recipient.islower() else recipient}.",
                         bold=True, space_after=6)
    add_paragraph_styled(doc, f"Về việc: {subject}", bold=True, space_after=12)

    if content:
        for paragraph in content.split("\n\n"):
            add_paragraph_styled(doc, paragraph.strip(), indent_cm=1.0, space_after=6)

    if bullets:
        for b in bullets:
            p = doc.add_paragraph(style="List Bullet")
            set_run_font(p.add_run(str(b)), size=BASE_SIZE)

    # Thời gian áp dụng
    if data.get("effective_from") or data.get("effective_to"):
        add_paragraph_styled(doc, "", space_after=0)
        if data.get("effective_from") and data.get("effective_to"):
            text = (f"Thời gian áp dụng: Từ {vn_date(data['effective_from'])} "
                    f"đến {vn_date(data['effective_to'])}.")
        elif data.get("effective_from"):
            text = f"Thời gian áp dụng: Từ {vn_date(data['effective_from'])}."
        else:
            text = f"Thời gian áp dụng: Đến {vn_date(data['effective_to'])}."
        add_paragraph_styled(doc, text, indent_cm=1.0, bold=True, space_after=6)

    add_paragraph_styled(
        doc,
        "Đề nghị toàn thể nhân viên chấp hành nghiêm túc. Mọi thắc mắc liên hệ trực tiếp "
        "với người ký thông báo.",
        indent_cm=1.0, space_after=18,
    )

    # Địa điểm + ngày ký
    place_line = f"{location + ', ' if location else ''}{vn_date(signed_date) if signed_date else ''}"
    if place_line.strip():
        p = doc.add_paragraph()
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(place_line)
        set_run_font(r, italic=True, size=BASE_SIZE)

    # Chữ ký (phía phải)
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(pt.add_run(signed_title.upper()), bold=True, size=BASE_SIZE)

    for _ in range(4):
        doc.add_paragraph()

    pname = doc.add_paragraph()
    pname.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(pname.add_run(signed_by), bold=True, size=BASE_SIZE)

    add_paragraph_styled(doc, f"({ten_quan})", align=WD_ALIGN_PARAGRAPH.RIGHT, italic=True,
                         size=BASE_SIZE - 1, space_after=0)

    add_footer(doc)
    return doc


def main():
    ap = argparse.ArgumentParser(description="Builder: Thông báo nội bộ.")
    ap.add_argument("--data", required=True)
    ap.add_argument("--config", default=None)
    ap.add_argument("--output", "-o", default="thong-bao.docx")
    args = ap.parse_args()

    if not Path(args.data).exists():
        print(f"Lỗi: không tìm thấy '{args.data}'.", file=sys.stderr); sys.exit(1)
    with open(args.data, "r", encoding="utf-8") as f:
        data = json.load(f)

    config = {}
    if args.config and Path(args.config).exists():
        with open(args.config, "r", encoding="utf-8") as f:
            config = yaml.safe_load(f) or {}

    if not data.get("subject"):
        print("Lỗi: thiếu 'subject'.", file=sys.stderr); sys.exit(2)

    doc = build_announcement(data, config)
    doc.save(args.output)
    print(f"OK: {args.output}")


if __name__ == "__main__":
    main()
