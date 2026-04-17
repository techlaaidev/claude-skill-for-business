"""
Builder: Biên bản bàn giao ca (ca sáng / chiều / tối).
Author: Techla — v1.0.0 — License: xem LICENSE.md

Usage:
    python scripts/builder_bien_ban_ca.py --data data.json --config config.yaml --output bien-ban.docx

data.json schema:
{
    "shift": "chiều",                     # sáng | chiều | tối
    "shift_date": "2026-04-17",
    "handover_by": "Nguyễn Văn A",
    "receive_by": "Trần Thị B",
    "cash_opening": 2000000,
    "cash_closing": 3800000,
    "materials": [                        # bảng nguyên liệu
        {"name": "Sữa tươi Vinamilk", "unit": "lít", "qty_start": 20, "qty_end": 8, "note": "cần đặt thêm"},
        {"name": "Cafe Robusta rang", "unit": "kg", "qty_start": 5, "qty_end": 3, "note": ""}
    ],
    "equipment": [                        # bảng thiết bị
        {"name": "Máy pha cafe La Marzocco", "status": "Bình thường", "note": ""},
        {"name": "Máy xay dao núi", "status": "Kêu to", "note": "cần vệ sinh"}
    ],
    "incidents": "Mất điện 30 phút lúc 15h, không ảnh hưởng lớn",
    "complaints": "1 khách phàn nàn matcha hơi đắng (đã pha lại)",
    "handover_notes": "Đặt thêm sữa tươi trong tối nay. Kiểm tra máy xay dao núi",
    "location": "Hà Nội"                  # optional
}
"""
from __future__ import annotations
import os
import sys
import json
import argparse
from pathlib import Path

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import yaml
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from doc_style import (
    new_doc, add_centered, add_paragraph_styled, add_footer,
    vn_date, set_run_font, BASE_SIZE
)
from vn_format import format_money


SHIFT_LABELS = {"sáng": "CA SÁNG", "chiều": "CA CHIỀU", "tối": "CA TỐI"}


def _shade(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _make_table(doc, headers, rows, widths_cm=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, size=BASE_SIZE - 1, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        _shade(hdr[i], "2E75B6")
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row in rows:
        tr = table.add_row().cells
        for i, val in enumerate(row):
            tr[i].text = ""
            p = tr[i].paragraphs[0]
            r = p.add_run(str(val))
            set_run_font(r, size=BASE_SIZE - 1)

    if widths_cm:
        for row in table.rows:
            for i, w in enumerate(widths_cm):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    doc.add_paragraph()


def build_handover(data: dict, config: dict) -> "Document":
    ten_quan = config.get("ten_quan", "Quán Cafe")
    dia_chi_quan = config.get("dia_chi_quan", "")

    shift = (data.get("shift") or "chiều").lower()
    shift_label = SHIFT_LABELS.get(shift, shift.upper())
    d = data["shift_date"]
    handover_by = data.get("handover_by", "(Chưa có)")
    receive_by = data.get("receive_by", "(Chưa có)")
    location = data.get("location", "")

    doc = new_doc()

    # Quốc hiệu rút gọn (biên bản không bắt buộc có)
    add_centered(doc, ten_quan.upper(), bold=True, size=14, space_after=0)
    if dia_chi_quan:
        add_centered(doc, dia_chi_quan, italic=True, size=11, space_after=6)
    add_centered(doc, "—————", space_after=6)

    add_centered(doc, "BIÊN BẢN BÀN GIAO CA", bold=True, size=16, space_before=6, space_after=0)
    add_centered(doc, shift_label, bold=True, italic=True, size=13, space_after=12)

    # Dòng địa điểm + ngày
    add_paragraph_styled(
        doc,
        f"Hôm nay, {vn_date(d)}{f', tại {location}' if location else ''}, "
        f"chúng tôi gồm:",
        indent_cm=1.0,
    )

    add_paragraph_styled(doc, f"Người giao ca: Ông/Bà {handover_by}", indent_cm=0.5, space_after=0)
    add_paragraph_styled(doc, f"Người nhận ca: Ông/Bà {receive_by}", indent_cm=0.5, space_after=12)

    add_paragraph_styled(doc, "Cùng tiến hành bàn giao ca với nội dung như sau:", indent_cm=1.0, space_after=12)

    # Phần 1 — Tiền quỹ
    add_paragraph_styled(doc, "1. TỒN QUỸ", bold=True, space_after=4)
    cash_open = data.get("cash_opening", 0) or 0
    cash_close = data.get("cash_closing", 0) or 0
    cash_diff = cash_close - cash_open
    add_paragraph_styled(doc, f"- Tồn quỹ đầu ca: {format_money(cash_open)}",
                         indent_cm=0.5, space_after=0)
    add_paragraph_styled(doc, f"- Tồn quỹ cuối ca: {format_money(cash_close)}",
                         indent_cm=0.5, space_after=0)
    sign = "+" if cash_diff >= 0 else ""
    add_paragraph_styled(doc, f"- Chênh lệch: {sign}{format_money(cash_diff)} "
                         f"(thu trong ca)", indent_cm=0.5, space_after=12)

    # Phần 2 — Tồn nguyên liệu
    add_paragraph_styled(doc, "2. TỒN NGUYÊN LIỆU", bold=True, space_after=4)
    materials = data.get("materials") or []
    if materials:
        rows = []
        for m in materials:
            rows.append([
                m.get("name", ""),
                m.get("unit", ""),
                str(m.get("qty_start", "")),
                str(m.get("qty_end", "")),
                m.get("note", ""),
            ])
        _make_table(doc, ["Nguyên liệu", "Đơn vị", "Tồn đầu ca", "Tồn cuối ca", "Ghi chú"], rows,
                    widths_cm=[5.5, 2, 2.5, 2.5, 4])
    else:
        add_paragraph_styled(doc, "(Không có dữ liệu)", indent_cm=0.5, italic=True, space_after=12)

    # Phần 3 — Thiết bị
    add_paragraph_styled(doc, "3. TÌNH TRẠNG THIẾT BỊ", bold=True, space_after=4)
    equipment = data.get("equipment") or []
    if equipment:
        rows = [[e.get("name", ""), e.get("status", ""), e.get("note", "")] for e in equipment]
        _make_table(doc, ["Thiết bị", "Tình trạng", "Ghi chú"], rows,
                    widths_cm=[7, 4, 5])
    else:
        add_paragraph_styled(doc, "(Không có dữ liệu)", indent_cm=0.5, italic=True, space_after=12)

    # Phần 4 — Sự cố
    add_paragraph_styled(doc, "4. SỰ CỐ TRONG CA", bold=True, space_after=4)
    add_paragraph_styled(doc, data.get("incidents") or "Không có sự cố.",
                         indent_cm=0.5, space_after=12)

    # Phần 5 — Khách phàn nàn
    add_paragraph_styled(doc, "5. Ý KIẾN KHÁCH HÀNG", bold=True, space_after=4)
    add_paragraph_styled(doc, data.get("complaints") or "Không có phàn nàn.",
                         indent_cm=0.5, space_after=12)

    # Phần 6 — Việc bàn giao
    add_paragraph_styled(doc, "6. VIỆC BÀN GIAO CHO CA SAU", bold=True, space_after=4)
    add_paragraph_styled(doc, data.get("handover_notes") or "Không có việc đặc biệt.",
                         indent_cm=0.5, space_after=18)

    # Cam kết
    add_paragraph_styled(
        doc,
        "Hai bên đã kiểm tra và xác nhận nội dung biên bản này là đúng với thực tế.",
        indent_cm=1.0, space_after=18,
    )

    # Chữ ký
    _signature_block(doc, handover_by, receive_by)

    add_footer(doc)
    return doc


def _signature_block(doc, handover_by, receive_by):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True

    def _sig(cell, title, name):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        set_run_font(r, size=BASE_SIZE, bold=True)
        for _ in range(4):
            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pb = cell.add_paragraph()
        pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rb = pb.add_run(f"(Ký, ghi rõ họ tên)\n{name}")
        set_run_font(rb, size=BASE_SIZE, bold=True)

    _sig(table.cell(0, 0), "NGƯỜI GIAO CA", handover_by)
    _sig(table.cell(0, 1), "NGƯỜI NHẬN CA", receive_by)


def main():
    ap = argparse.ArgumentParser(description="Builder: Biên bản bàn giao ca.")
    ap.add_argument("--data", required=True)
    ap.add_argument("--config", default=None)
    ap.add_argument("--output", "-o", default="bien-ban-ban-giao-ca.docx")
    args = ap.parse_args()

    if not Path(args.data).exists():
        print(f"Lỗi: không tìm thấy '{args.data}'.", file=sys.stderr); sys.exit(1)
    with open(args.data, "r", encoding="utf-8") as f:
        data = json.load(f)

    config = {}
    if args.config and Path(args.config).exists():
        with open(args.config, "r", encoding="utf-8") as f:
            config = yaml.safe_load(f) or {}

    required = ["shift_date"]
    missing = [f for f in required if not data.get(f)]
    if missing:
        print(f"Lỗi: thiếu field: {missing}", file=sys.stderr); sys.exit(2)

    doc = build_handover(data, config)
    doc.save(args.output)
    print(f"OK: {args.output}")


if __name__ == "__main__":
    main()
