"""
Style & helper chung cho 3 builder văn bản.
Author: Techla — v1.0.0 — License: xem LICENSE.md

Chuẩn:
- A4, margin 2cm
- Font Times New Roman 13pt (chuẩn hành chính VN)
- Tiêu đề đậm, căn giữa
- Footer "Soạn bởi Techla — Skill Soạn văn bản v1.0.0"
"""
from __future__ import annotations
from datetime import date as _date

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_FONT = "Times New Roman"
BASE_SIZE = 13
COLOR_GRAY = RGBColor(0x7F, 0x8C, 0x8D)

FOOTER_TEXT = "Soạn bởi Techla — Skill Soạn văn bản v1.0.0"


def set_run_font(run, *, size: int = BASE_SIZE, bold: bool = False,
                 italic: bool = False, color=None):
    run.font.name = BASE_FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), BASE_FONT)


def new_doc() -> Document:
    """Tạo doc mới đã setup A4 margin 2cm, font default."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    style = doc.styles["Normal"]
    style.font.name = BASE_FONT
    style.font.size = Pt(BASE_SIZE)
    return doc


def add_centered(doc: Document, text: str, *, size=BASE_SIZE, bold=False, italic=False,
                 color=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_paragraph_styled(doc: Document, text: str, *, size=BASE_SIZE, bold=False,
                         italic=False, indent_cm: float = 0, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if indent_cm:
        p.paragraph_format.first_line_indent = Cm(indent_cm)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_footer(doc: Document):
    """Footer cuối văn bản — in nghiêng, xám, căn giữa, 10pt."""
    doc.add_paragraph()
    add_centered(doc, FOOTER_TEXT, size=10, italic=True, color=COLOR_GRAY, space_before=12)


def vn_date(d) -> str:
    """'2026-05-01' hoặc date(2026,5,1) -> 'Ngày 01 tháng 05 năm 2026'."""
    if isinstance(d, str):
        d = _date.fromisoformat(d)
    return f"Ngày {d.day:02d} tháng {d.month:02d} năm {d.year}"


def vn_header_quoc_hieu(doc: Document):
    """Thêm quốc hiệu chuẩn 'CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM / Độc lập — Tự do — Hạnh phúc'."""
    add_centered(doc, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", bold=True, space_after=0)
    add_centered(doc, "Độc lập — Tự do — Hạnh phúc", bold=True, italic=True, space_after=6)
    # Dòng gạch dưới ngắn
    add_centered(doc, "—————", space_after=12)
