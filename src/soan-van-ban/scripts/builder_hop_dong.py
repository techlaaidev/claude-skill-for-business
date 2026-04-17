"""
Builder: Hợp đồng thử việc nhân viên pha chế / phục vụ quán cafe.
Author: Techla — v1.0.0 — License: xem LICENSE.md

Usage:
    python scripts/builder_hop_dong.py --data data.json --config config.yaml --output hop-dong.docx

data.json schema:
{
    "employee_name": "Nguyễn Văn A",
    "employee_dob": "1995-05-01",         # optional
    "employee_id_number": "025195000001", # CCCD, optional
    "employee_address": "...",            # optional
    "employee_phone": "...",              # optional
    "position": "Nhân viên pha chế",
    "salary_monthly": 8000000,            # OR salary_per_shift
    "salary_per_shift": null,
    "probation_months": 2,
    "start_date": "2026-05-01",
    "shift_description": "Ca sáng 7-13h / Ca chiều 13-19h",   # optional
    "allowance": "Ăn ca 30.000đ/ca",      # optional
    "signing_location": "Hà Nội",         # optional, default "Hà Nội"
    "signing_date": null                  # optional, null -> start_date
}
"""
from __future__ import annotations
import os
import sys
import json
import argparse
from pathlib import Path
from datetime import date

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import yaml
from docx.enum.text import WD_ALIGN_PARAGRAPH

from doc_style import (
    new_doc, add_centered, add_paragraph_styled, add_footer,
    vn_date, vn_header_quoc_hieu, set_run_font, BASE_SIZE
)
from vn_format import format_money, money_to_words_vi


def _load_json(path: str) -> dict:
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def _load_yaml(path: str) -> dict:
    with open(path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f) or {}


def build_contract(data: dict, config: dict) -> "Document":
    ten_quan = config.get("ten_quan", "Quán Cafe (chưa đặt tên)")
    dia_chi_quan = config.get("dia_chi_quan", "(chưa có địa chỉ)")
    ma_so_thue = config.get("ma_so_thue", "")
    nguoi_dai_dien = config.get("nguoi_dai_dien", "(Chủ quán)")
    chuc_vu_dai_dien = config.get("chuc_vu_dai_dien", "Chủ quán")

    name = data["employee_name"]
    position = data.get("position", "Nhân viên")
    start = data["start_date"]
    probation_months = int(data.get("probation_months", 2))
    salary_m = data.get("salary_monthly")
    salary_s = data.get("salary_per_shift")
    signing_loc = data.get("signing_location", "Hà Nội")
    signing_date = data.get("signing_date") or start

    doc = new_doc()
    vn_header_quoc_hieu(doc)

    add_centered(doc, "HỢP ĐỒNG THỬ VIỆC", bold=True, size=16, space_before=6, space_after=0)
    add_centered(doc, f"Số: HDTV-{start.replace('-', '')}-{name.split()[-1].upper() if name.split() else 'NV'}",
                 italic=True, space_after=12)

    add_paragraph_styled(
        doc,
        f"Hôm nay, {vn_date(signing_date)}, tại {signing_loc}, chúng tôi gồm:",
        indent_cm=1.0,
    )

    # Bên A
    add_paragraph_styled(doc, "BÊN A (NGƯỜI SỬ DỤNG LAO ĐỘNG):", bold=True, space_after=0)
    add_paragraph_styled(doc, f"Tên đơn vị: {ten_quan}", indent_cm=0.5, space_after=0)
    add_paragraph_styled(doc, f"Địa chỉ: {dia_chi_quan}", indent_cm=0.5, space_after=0)
    if ma_so_thue:
        add_paragraph_styled(doc, f"Mã số thuế: {ma_so_thue}", indent_cm=0.5, space_after=0)
    add_paragraph_styled(doc, f"Người đại diện: Ông/Bà {nguoi_dai_dien} — Chức vụ: {chuc_vu_dai_dien}",
                         indent_cm=0.5, space_after=12)

    # Bên B
    add_paragraph_styled(doc, "BÊN B (NGƯỜI LAO ĐỘNG):", bold=True, space_after=0)
    add_paragraph_styled(doc, f"Họ và tên: Ông/Bà {name}", indent_cm=0.5, space_after=0)
    if data.get("employee_dob"):
        add_paragraph_styled(doc, f"Ngày sinh: {vn_date(data['employee_dob'])}",
                             indent_cm=0.5, space_after=0)
    if data.get("employee_id_number"):
        add_paragraph_styled(doc, f"CCCD/CMND số: {data['employee_id_number']}",
                             indent_cm=0.5, space_after=0)
    if data.get("employee_address"):
        add_paragraph_styled(doc, f"Địa chỉ: {data['employee_address']}",
                             indent_cm=0.5, space_after=0)
    if data.get("employee_phone"):
        add_paragraph_styled(doc, f"Số điện thoại: {data['employee_phone']}",
                             indent_cm=0.5, space_after=0)
    add_paragraph_styled(doc, "", space_after=6)

    add_paragraph_styled(
        doc,
        "Sau khi thỏa thuận, hai bên cùng nhau ký kết Hợp đồng thử việc này với các điều khoản sau:",
        indent_cm=1.0,
        space_after=12,
    )

    # Điều 1 — Công việc và thời gian thử việc
    _dieu(doc, "Điều 1. Công việc và thời gian thử việc")
    add_paragraph_styled(doc, f"1.1. Vị trí công việc: {position}.", indent_cm=0.5, space_after=0)
    add_paragraph_styled(doc, f"1.2. Thời gian thử việc: {probation_months} tháng, "
                         f"kể từ {vn_date(start)}.", indent_cm=0.5, space_after=0)
    if data.get("shift_description"):
        add_paragraph_styled(doc, f"1.3. Giờ làm việc: {data['shift_description']}.",
                             indent_cm=0.5, space_after=0)
    add_paragraph_styled(doc, "", space_after=6)

    # Điều 2 — Lương và phụ cấp
    _dieu(doc, "Điều 2. Mức lương và phụ cấp")
    if salary_m:
        txt = f"2.1. Lương thử việc: {format_money(salary_m)} ({money_to_words_vi(salary_m)})/tháng."
        add_paragraph_styled(doc, txt, indent_cm=0.5, space_after=0)
    elif salary_s:
        txt = f"2.1. Lương thử việc: {format_money(salary_s)} ({money_to_words_vi(salary_s)})/ca."
        add_paragraph_styled(doc, txt, indent_cm=0.5, space_after=0)
    else:
        add_paragraph_styled(doc, "2.1. Lương thử việc: theo thỏa thuận.",
                             indent_cm=0.5, space_after=0)
    if data.get("allowance"):
        add_paragraph_styled(doc, f"2.2. Phụ cấp: {data['allowance']}.",
                             indent_cm=0.5, space_after=0)
    add_paragraph_styled(
        doc,
        "2.3. Lương trả một lần vào ngày cuối của mỗi tháng, qua tiền mặt hoặc chuyển khoản "
        "theo thỏa thuận.",
        indent_cm=0.5, space_after=6,
    )

    # Điều 3 — Quy định chung
    _dieu(doc, "Điều 3. Quy định chung về kỷ luật lao động")
    for line in [
        "3.1. Bên B phải đến làm việc đúng giờ, nghỉ làm phải báo trước ít nhất 24 giờ "
        "(trừ trường hợp bất khả kháng).",
        "3.2. Đi trễ quá 15 phút/lần hoặc nghỉ không phép sẽ bị trừ lương theo ca tương ứng, "
        "tái phạm nhiều lần sẽ bị chấm dứt thử việc.",
        "3.3. Bên B mặc đồng phục quán và tuân thủ quy định vệ sinh, thái độ phục vụ khách hàng.",
        "3.4. Không sử dụng điện thoại cá nhân trong giờ làm việc trừ khi có lý do công việc.",
    ]:
        add_paragraph_styled(doc, line, indent_cm=0.5, space_after=0)
    add_paragraph_styled(doc, "", space_after=6)

    # Điều 4 — Bảo mật
    _dieu(doc, "Điều 4. Bảo mật công thức và quy trình pha chế")
    for line in [
        "4.1. Bên B cam kết bảo mật tuyệt đối các công thức pha chế, quy trình kỹ thuật, "
        "công thức đồ uống độc quyền mà Bên A đã cung cấp hoặc đào tạo trong suốt quá trình làm việc "
        "và sau khi chấm dứt hợp đồng.",
        "4.2. Bên B không được chia sẻ, tiết lộ, sao chép dưới mọi hình thức (bao gồm chụp ảnh, quay "
        "video, ghi chép đem ra ngoài) các công thức và tài liệu đào tạo của Bên A.",
        "4.3. Trong vòng 6 tháng kể từ ngày chấm dứt hợp đồng, Bên B không được sử dụng các công thức "
        "này để kinh doanh cá nhân hoặc phục vụ cho đối thủ cạnh tranh trực tiếp của Bên A.",
        "4.4. Vi phạm điều khoản bảo mật sẽ phải bồi thường thiệt hại và chịu trách nhiệm theo quy định "
        "pháp luật.",
    ]:
        add_paragraph_styled(doc, line, indent_cm=0.5, space_after=0)
    add_paragraph_styled(doc, "", space_after=6)

    # Điều 5 — Chấm dứt
    _dieu(doc, "Điều 5. Chấm dứt hợp đồng thử việc")
    for line in [
        "5.1. Trong thời gian thử việc, mỗi bên có quyền đơn phương chấm dứt hợp đồng bằng thông báo "
        "bằng văn bản hoặc tin nhắn trước 03 ngày làm việc, không cần bồi thường.",
        "5.2. Hết thời gian thử việc, nếu Bên B đạt yêu cầu, hai bên sẽ ký hợp đồng lao động chính thức. "
        "Nếu không đạt, Bên A thông báo bằng văn bản và thanh toán đầy đủ lương đã làm việc.",
        "5.3. Bên A có quyền chấm dứt ngay không cần báo trước nếu Bên B: trộm cắp, đánh nhau trong "
        "quán, vi phạm nghiêm trọng bảo mật công thức, hoặc gây thiệt hại tài sản lớn cho quán.",
    ]:
        add_paragraph_styled(doc, line, indent_cm=0.5, space_after=0)
    add_paragraph_styled(doc, "", space_after=6)

    # Điều 6 — Điều khoản chung
    _dieu(doc, "Điều 6. Điều khoản chung")
    for line in [
        "6.1. Hợp đồng này được lập thành 02 bản có giá trị pháp lý như nhau, mỗi bên giữ 01 bản.",
        "6.2. Mọi sửa đổi, bổ sung phải được hai bên thỏa thuận và lập thành văn bản.",
        "6.3. Tranh chấp phát sinh (nếu có) được giải quyết trên cơ sở thương lượng. Không thương lượng "
        "được, hai bên có quyền yêu cầu cơ quan có thẩm quyền giải quyết.",
    ]:
        add_paragraph_styled(doc, line, indent_cm=0.5, space_after=0)
    add_paragraph_styled(doc, "", space_after=18)

    # Chữ ký
    _signature_block(doc, ten_quan, nguoi_dai_dien, name)

    add_footer(doc)
    return doc


def _dieu(doc, title: str):
    add_paragraph_styled(doc, title, bold=True, space_after=4)


def _signature_block(doc, ten_quan, nguoi_dai_dien, employee_name):
    # Bảng 2 cột cho chữ ký
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    left = table.cell(0, 0)
    right = table.cell(0, 1)

    def _sig(cell, label_top, label_bottom):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label_top)
        set_run_font(r, size=BASE_SIZE, bold=True)
        for _ in range(4):
            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pb = cell.add_paragraph()
        pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rb = pb.add_run(label_bottom)
        set_run_font(rb, size=BASE_SIZE, bold=True)

    _sig(left, "ĐẠI DIỆN BÊN A", f"(Ký, ghi rõ họ tên)\n{nguoi_dai_dien}")
    _sig(right, "BÊN B", f"(Ký, ghi rõ họ tên)\n{employee_name}")


def main():
    ap = argparse.ArgumentParser(description="Builder: Hợp đồng thử việc.")
    ap.add_argument("--data", required=True, help="JSON file chứa thông tin nhân viên + hợp đồng")
    ap.add_argument("--config", default=None, help="YAML file chứa thông tin quán")
    ap.add_argument("--output", "-o", default="hop-dong-thu-viec.docx")
    args = ap.parse_args()

    if not Path(args.data).exists():
        print(f"Lỗi: không tìm thấy '{args.data}'.", file=sys.stderr); sys.exit(1)
    data = _load_json(args.data)

    config = {}
    if args.config and Path(args.config).exists():
        config = _load_yaml(args.config)

    required = ["employee_name", "position", "start_date"]
    missing = [f for f in required if not data.get(f)]
    if missing:
        print(f"Lỗi: thiếu field bắt buộc: {missing}", file=sys.stderr); sys.exit(2)

    doc = build_contract(data, config)
    doc.save(args.output)
    print(f"OK: {args.output}")


if __name__ == "__main__":
    main()
